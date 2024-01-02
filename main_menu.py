import psycopg2
from PyQt5.QtCore import QEvent, Qt
from PyQt5.QtWidgets import QWidget, QPushButton, QVBoxLayout, QStackedWidget, QLabel, QHBoxLayout, QTextBrowser, \
    QListWidget, QListWidgetItem, QMessageBox, QShortcut
from PyQt5.QtGui import QKeySequence

from docx import Document
from datetime import datetime, timedelta

from conn_db import connect, close_db_connect
from listeners_page import ListenersPage
from teachers_page import TeachersPage
from course_page import CoursesPage
from contracts_page import ContractsPage
from reports_page import ReportsPage


class MainMenuWidget(QWidget):
    def __init__(self, parent):
        super().__init__()
        self.parent = parent
        self.initUI()
        self.load_requests()

    def initUI(self):
        self.is_request_selected = False  # Флаг для отслеживания выбора заявки

        self.taskbar_layout = QVBoxLayout()

        self.pages = QStackedWidget()

        page_names = ["Главное меню", "Слушатели", "Курсы", "Преподаватели", "Договоры", "Отчеты"]
        for i, page_name in enumerate(page_names):
            button = QPushButton(page_name)
            button.clicked.connect(self.show_page(i))
            self.taskbar_layout.addWidget(button)

        self.create_main_menu_page()
        self.create_listeners_page()
        self.create_course_page()
        self.create_teachers_page()
        self.create_contracts_page()
        self.create_reports_page()

        self.requests_list.itemClicked.connect(self.show_request_details)

        # Размещаем виджеты с кнопками и страницами в основном макете
        main_layout = QHBoxLayout(self)
        main_layout.addLayout(self.taskbar_layout)
        main_layout.addWidget(self.pages)

        reload_shortcut = QShortcut(QKeySequence(Qt.Key_F5), self)
        reload_shortcut.activated.connect(self.load_requests)

        self.setLayout(main_layout)

        self.installEventFilter(self)

    def eventFilter(self, obj, event):
        if obj == self and event.type() == QEvent.MouseButtonRelease:
            self.clear_listener_selection()
        return super().eventFilter(obj, event)

    def clear_listener_selection(self):
        # Очистите выделение в списке слушателей
        self.requests_list.clearSelection()
        self.details_text.clear()
        self.is_request_selected = False

    def resizeEvent(self, event):
        # При изменении размера окна будем изменять размер шрифта
        new_font_size = int(self.width() / 60)  # Пример расчета нового размера шрифта
        font = self.details_text.font()
        font.setPointSize(new_font_size)
        self.details_text.setFont(font)

        # Вызываем родительский метод для обработки стандартных событий изменения размера окна
        super().resizeEvent(event)

    def create_main_menu_page(self):
        page = QWidget()
        layout = QHBoxLayout()

        # Область с заявками
        requests_area = QWidget()
        requests_layout = QVBoxLayout()
        label = QLabel("Заявки:")
        self.requests_list = QListWidget()
        requests_layout.addWidget(label)
        requests_layout.addWidget(self.requests_list)
        requests_area.setLayout(requests_layout)

        # Область с подробной информацией
        details_area = QWidget()
        details_layout = QVBoxLayout()
        label = QLabel("Подробная информация о заявке:")
        self.details_text = QTextBrowser()
        details_layout.addWidget(label)
        details_layout.addWidget(self.details_text)
        details_area.setLayout(details_layout)

        buttons_area = QWidget()
        buttons_layout = QVBoxLayout()
        compose_contract_button = QPushButton("Составить договор")
        compose_contract_button.clicked.connect(self.generate_contract)
        reject_button = QPushButton("Отказать")
        reject_button.clicked.connect(self.reject_request)
        buttons_layout.addWidget(compose_contract_button)
        buttons_layout.addWidget(reject_button)
        buttons_area.setLayout(buttons_layout)

        self.is_request_selected = True

        layout.addWidget(requests_area)
        layout.addWidget(details_area)
        layout.addWidget(buttons_area)

        page.setLayout(layout)

        self.pages.addWidget(page)

    def update_requests(self):
        self.load_requests()

    def showEvent(self, event):
        # Вызывается при каждом показе страницы
        self.update_requests()

    def show_page(self, page_index):
        def handler():
            self.pages.setCurrentIndex(page_index)

        return handler

    def load_requests(self):
        connection = connect()

        cursor = connection.cursor()

        cursor.execute(
            'SELECT r.request_id, r.first_name, r.second_name, r.patronymic, r.phone_number, r.age, r.status, '
            'r.course_id '
            'FROM "Requests" r '
            'WHERE r.status = FALSE '
            'ORDER BY r.request_id ASC'
        )

        requests = cursor.fetchall()

        close_db_connect(connection, cursor)

        self.requests_list.clear()

        for request in requests:
            request_id, first_name, second_name, patronymic, phone_number, age, course_id, status = request
            item_text = f"Заявка № {request_id}"
            item = QListWidgetItem(item_text)
            item.setData(1, request_id)
            self.requests_list.addItem(item)

    def show_request_details(self, item):
        request_id = item.data(1)

        self.is_request_selected = True

        connection = connect()

        cursor = connection.cursor()
        cursor.execute(
            'SELECT r.request_id, r.first_name, r.second_name, r.age, r.status, r.patronymic, r.phone_number, '
            'c.name_course '
            'FROM "Requests" r '
            'JOIN "Courses" c ON r.course_id = c.course_id '
            'WHERE r.request_id = %s',
            (request_id,)
        )
        request_details = cursor.fetchone()

        close_db_connect(connection, cursor)

        if request_details:
            request_id, first_name, second_name, age, status, patronymic, phone_number, course_name = request_details
            details_text = f"Имя: {first_name}\nФамилия: {second_name}\nОтчество: {patronymic}\n" \
                           f"Возраст: {age}\nТелефон: {phone_number}\nКурс: {course_name}\n"

            self.details_text.setPlainText(details_text)

    def generate_contract(self):
        # Проверка, выбрана ли заявка
        if not self.is_request_selected:
            QMessageBox.warning(self, 'Предупреждение', 'Выберите заявку перед составлением договора')
            return
        month_translation = {
            'January': 'января',
            'February': 'февраля',
            'March': 'марта',
            'April': 'апреля',
            'May': 'мая',
            'June': 'июня',
            'July': 'июля',
            'August': 'августа',
            'September': 'сентября',
            'October': 'октября',
            'November': 'ноября',
            'December': 'декабря'
        }

        template_path = 'C:/Users/Александр/PycharmProjects/Course_app/основные файлы/Договор.docx'
        doc = Document(template_path)

        start_date = datetime.now() + timedelta(days=7)

        end_date = datetime.now() + timedelta(days=180)

        connection = connect()
        cursor = connection.cursor()

        try:
            request_id = self.requests_list.currentItem().data(1)
            cursor.execute(
                'SELECT r.request_id, r.first_name, r.second_name, r.age, r.status, r.patronymic, r.phone_number, '
                'r.course_id, c.name_course '
                'FROM "Requests" r '
                'JOIN "Courses" c ON r.course_id = c.course_id '
                'WHERE r.request_id = %s',
                (request_id,)
            )
            request_details = cursor.fetchone()
            cursor.execute('SELECT MAX(contract_id) FROM "Contracts"')
            last_contract_id = cursor.fetchone()[0]

            if last_contract_id is None:
                new_contract_id = 1
            else:
                new_contract_id = last_contract_id + 1

            cursor.execute(
                'INSERT INTO "Contracts" (contract_id, contract_date, contract_file, user_id, request_id, signed) '
                'VALUES (%s, %s, %s, %s, %s, %s)',
                (new_contract_id, datetime.now(), None, None, request_id, False)
            )

            if request_details:
                request_id, first_name, second_name, age, status, patronymic, phone_number, course_id, course_name =\
                    request_details

                # Замена меток в документе на соответствующие данные
                for paragraph in doc.paragraphs:
                    for run in paragraph.runs:
                        text = run.text
                        if 'date' in text:
                            run.text = text.replace('date', datetime.now().strftime('%d'))
                        elif 'month' in text:
                            rus_month = month_translation[datetime.now().strftime('%B')]
                            run.text = text.replace('month', rus_month)
                        elif 'year' in text:
                            run.text = text.replace('year', datetime.now().strftime('%Y'))
                        elif 'start' in text:
                            run.text = text.replace('start', start_date.strftime('%d'))
                        elif 'period' in text:
                            rus_period = month_translation[start_date.strftime('%B')]
                            run.text = text.replace('period', rus_period)
                        elif 'annual' in text:
                            run.text = text.replace('annual', start_date.strftime('%Y'))
                        elif 'end' in text:
                            run.text = text.replace('end', end_date.strftime('%d'))
                        elif 'perm' in text:
                            rus_perm = month_translation[end_date.strftime('%B')]
                            run.text = text.replace('perm', rus_perm)
                        elif 'god' in text:
                            run.text = text.replace('god', end_date.strftime('%Y'))
                        elif 'numberdog' in text:
                            run.text = text.replace('numberdog', str(new_contract_id))
                        elif 'contract' in text:
                            run.text = text.replace('contract', str(new_contract_id))
                        elif 'FIOuser' in text:
                            run.text = text.replace('FIOuser', f'{second_name} {first_name} {patronymic}')
                        elif 'userphone' in text:
                            run.text = text.replace('userphone', phone_number)
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    text = run.text
                                    if 'FIOuser' in text:
                                        run.text = text.replace('FIOuser', f'{second_name} {first_name} {patronymic}')
                                    elif 'userphone' in text:
                                        run.text = text.replace('userphone', phone_number)
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            cell.text = cell.text.replace('coursename', course_name)

            output_path = f'C:/Users/Александр/PycharmProjects/Course_app/основные файлы/Договор_{new_contract_id}.docx'

            doc.save(output_path)

            with open(output_path, 'rb') as file:
                file_data = file.read()
            cursor.execute(
                'UPDATE "Contracts" SET contract_file = %s WHERE contract_id = %s',
                (psycopg2.Binary(file_data), new_contract_id)
            )

            connection.commit()

            print(f"Договор успешно сохранен в базе данных с номером {new_contract_id}")
            self.clear_listener_selection()

            cursor.execute('UPDATE "Requests" SET status = TRUE WHERE request_id = %s', (request_id,))
            connection.commit()

            print(f"Статус заявки с ID {request_id} успешно обновлен")
            self.update_requests()

        except Exception as e:
            print(f"Ошибка при сохранении договора в базу данных: {e}")
            connection.rollback()

        finally:
            close_db_connect(connection, cursor)

    def reject_request(self):
        if not self.is_request_selected:
            QMessageBox.warning(self, 'Предупреждение', 'Выберите заявку перед отказом')
            return

        try:
            request_id = self.requests_list.currentItem().data(1)

            connection = connect()
            cursor = connection.cursor()
            cursor.execute('DELETE FROM "Requests" WHERE request_id = %s', (request_id,))
            connection.commit()
            print(f"Заявка с ID {request_id} успешно удалена")

            self.clear_listener_selection()
            self.update_requests()

        except Exception as e:
            print(f"Ошибка при удалении заявки из базы данных: {e}")
            connection.rollback()

        finally:
            close_db_connect(connection, cursor)

    def create_listeners_page(self):
        listeners_page = ListenersPage(self)
        self.pages.addWidget(listeners_page)

    def create_teachers_page(self):
        teachers_page = TeachersPage(self)
        self.pages.addWidget(teachers_page)

    def create_course_page(self):
        course_page = CoursesPage(self)
        self.pages.addWidget(course_page)

    def create_contracts_page(self):
        contracts_page = ContractsPage(self)
        self.pages.addWidget(contracts_page)

    def create_reports_page(self):
        reports_page = ReportsPage(self)
        self.pages.addWidget(reports_page)
